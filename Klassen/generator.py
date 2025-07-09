# -*- coding: utf-8 -*-
"""
Dieses Modul definiert den DocxGenerator.

Die Klasse ist verantwortlich für die Erstellung des finalen Word-Katalogs
basierend auf den aufbereiteten Daten und den Konfigurationen.
"""

import datetime
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, RGBColor
from PIL import Image


class DocxGenerator:
    """Erstellt das Word-Dokument für den Ersatzteilkatalog."""

    def __init__(
        self,
        data: dict,
        main_bom,
        author_name: str,
        custom_doc_number: str,
        template_path: str,
        output_path: str,
        auto_update_fields: bool,
        project_path: str,
        config_manager,
    ):
        self.data = data
        self.main_bom = main_bom
        self.author_name = author_name
        self.custom_doc_number = custom_doc_number
        self.template_path = template_path
        self.output_path = output_path
        self.auto_update_fields = auto_update_fields
        self.project_path = project_path
        self.config = config_manager
        self.doc = Document(self.template_path)

    def run(self) -> bool:
        """Führt den gesamten Prozess der Dokumenterstellung aus."""
        try:
            formatting_options = self.config.config.get("formatting_options", {})
            
            self._update_header_footer()
            self._create_toc(formatting_options)
            self._create_assembly_section(self.data)
            self._replace_graphic_placeholders()
            self._insert_blank_pages(formatting_options)
            self._create_cover_sheet(formatting_options)

            self.doc.save(self.output_path)

            if self.auto_update_fields:
                self._update_fields_and_refs_with_word()

            return True
        except Exception as e:
            print(f"FEHLER bei der Dokumenterstellung: {e}")
            return False

    def _update_fields_and_refs_with_word(self):
        """
        Startet Word, aktualisiert das Inhaltsverzeichnis und ersetzt die
        Seiten-Referenzen.
        """
        print("Starte erweiterte Feld-Aktualisierung mit MS Word...")
        if sys.platform != "win32":
            print("WARNUNG: Automatische Aktualisierung ist nur auf Windows möglich.")
            return

        word = None
        try:
            import win32com.client as win32

            abs_path = os.path.abspath(self.output_path)
            word = win32.Dispatch("Word.Application")
            word.Visible = False

            # Das Dokument wird jetzt direkt geöffnet. Der "Säuberungs"-Schritt
            # ist für .docx-Dateien nicht notwendig.
            doc = word.Documents.Open(abs_path)

            print("  -> Aktualisiere alle Dokumenten-Felder (z.B. Seitenzahlen)...")
            doc.Fields.Update()

            print("  -> Aktualisiere Inhaltsverzeichnis...")
            if doc.TablesOfContents.Count > 0:
                doc.TablesOfContents(1).Update()

            page_map = {}
            print("  -> Lese Seitenzahlen aus Inhaltsverzeichnis-Text...")
            if doc.TablesOfContents.Count > 0:
                toc_text = doc.TablesOfContents(1).Range.Text
                pattern = re.compile(
                    r"([A-Z0-9\.\- ]+-[A-Z0-9\.\- ]+).*?\t(\d+)", re.IGNORECASE
                )
                for match in pattern.finditer(toc_text):
                    key = match.group(1).strip()
                    key = key.split('(')[0].strip()
                    page_map[key] = match.group(2)
            
            if not page_map and doc.TablesOfContents.Count > 0:
                print("    [WARNUNG] Konnte keine Seitenzahlen aus dem Inhaltsverzeichnis extrahieren.")

            if page_map:
                print("  -> Ersetze Seiten-Referenzen im Dokument...")
                for table in doc.Tables:
                    for row in table.Rows:
                        for cell in row.Cells:
                            match = re.search(r"\[REF:(.*?)\]", cell.Range.Text)
                            if match and match.group(1) in page_map:
                                cell.Range.Text = f"S. {page_map[match.group(1)]}"

            doc.Close(SaveChanges=True)
            word.Quit()
            word = None
            print("Erweiterte Feld-Aktualisierung erfolgreich abgeschlossen.")

        except ImportError:
            print("FEHLER: Die Bibliothek 'pywin32' wird für die automatische Aktualisierung benötigt.")
        except Exception as e:
            print(f"FEHLER bei der automatischen Aktualisierung: {e}")
        finally:
            if word is not None:
                word.Quit()

    def _create_table_for_assembly(self, items):
        """Erstellt die Tabelle dynamisch basierend auf der Konfiguration."""
        output_columns = self.config.config.get("output_columns", [])
        if not output_columns:
            print("WARNUNG: Keine Ausgabespalten in der Konfiguration definiert.")
            return

        table_styles = self.config.config.get("table_styles", {})
        base_style = table_styles.get("base_style", "Table Grid")
        header_bold = table_styles.get("header_bold", True)
        header_font_color_hex = table_styles.get("header_font_color")
        header_shading_color = table_styles.get("header_shading_color", "4F81BD")
        shading_enabled = table_styles.get("shading_enabled", True)
        shading_color = table_styles.get("shading_color", "DAE9F8")

        table = self.doc.add_table(rows=1, cols=len(output_columns))
        table.style = base_style
        table.width = Cm(16.5)
        table.autofit = False
        table.allow_autofit = False

        hdr_cells = table.rows[0].cells
        tr = table.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement("w:tblHeader")
        trPr.append(tblHeader)
        for i, col_config in enumerate(output_columns):
            cell = hdr_cells[i]
            cell.text = col_config.get("header", "")
            
            run = cell.paragraphs[0].runs[0]
            if header_bold:
                run.font.bold = True
            if header_font_color_hex:
                try:
                    run.font.color.rgb = RGBColor.from_string(header_font_color_hex)
                except ValueError:
                    print(f"WARNUNG: Ungültiger Hex-Code für Header-Schriftfarbe: '{header_font_color_hex}'")

            if header_shading_color:
                self._set_cell_shading(cell, header_shading_color)

        sorted_items = sorted(
            items,
            key=lambda item: (
                float(item.get("POS", 0))
                if str(item.get("POS", "0")).replace(".", "", 1).isdigit()
                else float("inf")
            ),
        )

        for i, item in enumerate(sorted_items):
            row_cells = table.add_row().cells
            for col_idx, col_config in enumerate(output_columns):
                source_id = col_config.get("source_id")
                col_id = col_config.get("id")
                key_for_data = source_id if source_id else col_id

                cell_text = ""
                if col_id == "std_seite":
                    if item.get("is_assembly", False) and item.get("Teilenummer"):
                        teilenummer_raw = str(item.get("Teilenummer", ""))
                        no_newlines = teilenummer_raw.replace('\n', '').replace('\r', '')
                        clean_znr = no_newlines.split("(")[0].strip().replace(" ", "")
                        if clean_znr:
                            cell_text = f"[REF:{clean_znr}]"
                else:
                    value = item.get(key_for_data, "")
                    if key_for_data == "POS" and isinstance(value, (int, float)):
                        cell_text = f"{value:g}"
                    else:
                        cell_text = str(value)

                self._add_multiline_text(row_cells[col_idx], cell_text)

            if shading_enabled and (i % 2) != 0:
                for cell in row_cells:
                    self._set_cell_shading(cell, shading_color)

        total_percent = sum(c.get("width_percent", 10) for c in output_columns)
        if total_percent > 0:
            for i, col_config in enumerate(output_columns):
                percent = col_config.get("width_percent", 10)
                col_width = int(
                    table.width.cm * (percent / total_percent) * 360000
                )
                for row in table.rows:
                    row.cells[i].width = col_width
                table.columns[i].width = col_width

    def _create_cover_sheet(self, formatting_options: dict):
        """Erstellt das Deckblatt, entweder Standard oder aus externer Datei."""
        cover_type = formatting_options.get("cover_sheet_type", "default")
        cover_path = formatting_options.get("cover_sheet_path", "")

        if cover_type == "external_docx" and os.path.exists(cover_path):
            print(f"INFO: Füge externes Deckblatt ein aus: {cover_path}")
            self._insert_docx_content(cover_path)

            p = self.doc.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)
        else:
            self._create_default_cover_sheet()
        
        p = self.doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
        self.doc.element.body.insert(0, p._p)
        self.doc.element.body.remove(p._p)

    def _insert_blank_pages(self, formatting_options: dict):
        """Fügt leere oder externe Seiten vor dem Inhaltsverzeichnis ein."""
        pages_type = formatting_options.get("blank_pages_type", "blank")
        pages_path = formatting_options.get("blank_pages_path", "")
        num_pages = formatting_options.get("blank_pages_before_toc", 0)

        if pages_type == "external_docx" and os.path.exists(pages_path):
            print(f"INFO: Füge externe Seite(n) ein aus: {pages_path}")
            self._insert_docx_content(pages_path)
        else:
            # Füge leere Seiten durch Seitenumbrüche am Anfang ein
            for _ in range(num_pages):
                p_element = OxmlElement("w:p")
                r_element = OxmlElement("w:r")
                br_element = OxmlElement("w:br")
                br_element.set(qn("w:type"), "page")
                r_element.append(br_element)
                p_element.append(r_element)
                self.doc.element.body.insert(0, p_element)

    def _create_default_cover_sheet(self):
        elements = []
        temp_doc = Document()
        
        p_title = temp_doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run(self.main_bom.titel or "[TITEL]")
        font_title = run_title.font; font_title.name = 'Calibri'; font_title.size = Cm(1.5); font_title.bold = True
        elements.append(p_title._p)
        
        p_line = temp_doc.add_paragraph("_________________________________________________")
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elements.append(p_line._p)
        
        p_subject = temp_doc.add_paragraph()
        p_subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_subject = p_subject.add_run("Ersatzteilkatalog"); run_subject.font.size = Cm(0.8)
        elements.append(p_subject._p)
        
        p_grafik = temp_doc.add_paragraph()
        p_grafik.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_path = next((os.path.join(self.project_path, "Grafik", f"EL{ext}") for ext in ['.png', '.jpg'] if os.path.exists(os.path.join(self.project_path, "Grafik", f"EL{ext}"))), None)
        if image_path:
            self._add_scaled_picture(p_grafik, image_path, Cm(16), Cm(15))
        elements.append(p_grafik._p)

        # Füge die Elemente in umgekehrter Reihenfolge am Anfang des Hauptdokuments ein
        for element in reversed(elements):
            self.doc.element.body.insert(0, element)

    def _create_toc(self, formatting_options: dict):
        if formatting_options.get("toc_on_new_page", True):
            self.doc.add_page_break()
        self.doc.add_heading("Inhaltsverzeichnis", level=1)
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fldChar_begin = OxmlElement("w:fldChar"); fldChar_begin.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve"); instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        fldChar_end = OxmlElement("w:fldChar"); fldChar_end.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar_begin); run._r.append(instrText); run._r.append(fldChar_end)
        self.doc.add_page_break()

    def _add_scaled_picture(self, paragraph, image_path, max_width, max_height):
        try:
            with Image.open(image_path) as img:
                original_width_px, original_height_px = img.size
            if original_width_px == 0 or original_height_px == 0: return
            
            aspect_ratio = float(original_height_px) / float(original_width_px)
            if max_width * aspect_ratio > max_height:
                paragraph.add_run().add_picture(image_path, height=max_height)
            else:
                paragraph.add_run().add_picture(image_path, width=max_width)
        except Exception as e:
            print(f"Bildfehler: {e}")
            paragraph.add_run(f"[Bild {os.path.basename(image_path)} nicht ladbar]").italic = True

    def _replace_graphic_placeholders(self):
        grafik_folder = os.path.join(self.project_path, "Grafik")
        for p in self.doc.paragraphs:
            if "[GRAFIK_PLATZHALTER_" in p.text:
                znr = p.text.split('_')[-1].replace(']', '')
                image_path = next((os.path.join(grafik_folder, f"{znr}{ext}") for ext in ['.png', '.jpg', '.jpeg'] if os.path.exists(os.path.join(grafik_folder, f"{znr}{ext}"))), None)
                p.clear() 
                if image_path:
                    self._add_scaled_picture(p, image_path, Cm(16), Cm(19))
                else:
                    p.add_run("[Keine Grafik zugeordnet]").italic = True

    def _update_header_footer(self):
        if self.custom_doc_number:
            full_zeich_nr = self.custom_doc_number
        else:
            full_zeich_nr = f"{self.main_bom.kundennummer or self.main_bom.zeichnungsnummer or ''} (EL)"
            
        full_title = f"{self.main_bom.titel or ''} - {self.main_bom.zusatzbenennung or ''}".strip(' -')
        creation_date = datetime.date.today().strftime('%d.%m.%Y')
        replacements = {'[TITEL]': full_title, '[THEMA]': "Ersatzteilkatalog", '[ZEICH]': full_zeich_nr, '[VERWEND]': self.main_bom.verwendung or "", '[AUTOR]': self.author_name, '[EDATUM]': creation_date}
        
        for section in self.doc.sections:
            for part in [section.header, section.footer, section.first_page_header, section.first_page_footer]:
                if part:
                    for table in part.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    for run in p.runs:
                                        for key, value in replacements.items():
                                            if key in run.text: run.text = run.text.replace(key, str(value))
                    for p in part.paragraphs: 
                        for run in p.runs:
                            for key, value in replacements.items():
                                if key in run.text: run.text = run.text.replace(key, str(value))

    def _create_assembly_section(self, assembly_data):
        if not assembly_data:
            return
        formatting_options = self.config.config.get("formatting_options", {})
        title_format = formatting_options.get(
            "assembly_title_format", "{benennung} ({teilenummer})"
        )
        
        add_space = formatting_options.get("add_space_after_title", True)
        table_on_new_page = formatting_options.get("table_on_new_page", False)
        
        title = title_format.format(
            benennung=assembly_data.get('Benennung', ''),
            teilenummer=assembly_data.get('Teilenummer', '')
        )
        
        heading = self.doc.add_heading(title, level=1)
        heading.paragraph_format.keep_with_next = True
        
        styles = self.doc.styles
        if 'Überschrift 1' in styles:
            heading.style = styles['Überschrift 1']
        elif 'Heading 1' in styles:
            heading.style = styles['Heading 1']

        if add_space:
            self.doc.add_paragraph()

        p_grafik = self.doc.add_paragraph()
        p_grafik.add_run(f"[GRAFIK_PLATZHALTER_{assembly_data.get('Teilenummer')}]")
        p_grafik.paragraph_format.keep_with_next = True
        if table_on_new_page:
            self.doc.add_page_break()
        else:
            self.doc.add_paragraph()
        
        if assembly_data.get('children'):
            self._create_table_for_assembly(assembly_data.get('children'))
        
        self.doc.add_page_break()
        for child in assembly_data.get('children', []):
            if child.get('is_assembly'):
                self._create_assembly_section(child)

    def _add_multiline_text(self, cell, text):
        cell.text = ''
        p = cell.paragraphs[0]
        for i, line in enumerate(str(text).split('\n')):
            if i > 0: p.add_run().add_break()
            p.add_run(line)
            
    def _set_cell_shading(self, cell, fill_color):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), fill_color)
        tc_pr.append(shd)

    def _insert_docx_content(self, docx_path):
        """Fügt den Inhalt eines anderen DOCX-Dokuments am Anfang des Hauptdokuments ein."""
        try:
            source_doc = Document(docx_path)
            for element in reversed(source_doc.element.body):
                if element.tag.endswith('sectPr'):
                    continue
                self.doc.element.body.insert(0, element)
        except Exception as e:
            print(f"FEHLER beim Einfügen von '{docx_path}': {e}")
            p = self.doc.add_paragraph(f"[Fehler: Inhalt aus '{os.path.basename(docx_path)}' konnte nicht geladen werden.]")
            # Diesen temporären Paragraphen an den Anfang verschieben und den am Ende erstellten löschen
            self.doc.element.body.insert(0, p._p)
            self.doc.element.body.remove(p._p)